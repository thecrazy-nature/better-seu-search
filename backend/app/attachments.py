from __future__ import annotations

import io
import re
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET
from urllib.parse import urlparse


MAX_ATTACHMENT_TEXT_CHARS = 12000
MAX_ATTACHMENT_PART_CHARS = 4000
MAX_ATTACHMENT_TABLE_ROWS = 80
MAX_PDF_TABLE_ROWS_PER_ATTACHMENT = 60
MAX_PDF_TABLE_ROWS_PER_PAGE = 12
MAX_TABLE_ROW_TEXT_CHARS = 900
MIN_LEGACY_TEXT_SCORE = 18
KNOWN_ATTACHMENT_EXTENSIONS = {
    ".pdf",
    ".doc",
    ".docx",
    ".xls",
    ".xlsx",
    ".zip",
    ".rar",
    ".jpg",
    ".jpeg",
    ".png",
}
WORD_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
SHEET_NS = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"
REL_NS = "{http://schemas.openxmlformats.org/package/2006/relationships}"
OFFICE_REL_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
OLE2_SIGNATURE = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
LEGACY_COMMON_TERMS = {
    "东南大学",
    "附件",
    "申请",
    "学院",
    "学生",
    "本科",
    "研究生",
    "课程",
    "教务",
    "专业",
    "项目",
    "通知",
    "报名",
    "考试",
    "成绩",
    "姓名",
    "学号",
    "一卡通",
    "联系电话",
    "序号",
    "填表",
    "材料",
    "时间",
    "地点",
}
LEGACY_OFFICE_NOISE_TERMS = {
    "Root Entry",
    "SummaryInformation",
    "DocumentSummaryInformation",
    "WordDocument",
    "Workbook",
    "ETExtData",
    "WPS Office",
    "Normal dotm",
    "Arial",
    "Calibri",
    "Times New Roman",
    "Microsoft Office Word",
    "Microsoft YaHei UI",
    "TableStyle",
    "PivotStyle",
    "PAGEREF",
    "Content_Types",
    "theme/theme",
    "themeManager",
}


def attachment_extension(url: str, name: str = "") -> str:
    url_suffix = Path(urlparse(url).path).suffix.lower()
    name_suffix = Path(name).suffix.lower()
    if url_suffix in KNOWN_ATTACHMENT_EXTENSIONS:
        return url_suffix
    return name_suffix or url_suffix


def infer_attachment_extension(content: bytes, url: str, name: str = "", content_type: str = "") -> str:
    ext = attachment_extension(url, name)
    lowered_type = (content_type or "").lower()
    head = content[:8]
    if head.startswith(b"%PDF"):
        return ".pdf"
    if head.startswith(b"PK\x03\x04"):
        if _zip_has(content, "word/document.xml"):
            return ".docx"
        if _zip_has(content, "xl/workbook.xml"):
            return ".xlsx"
        return ext
    if head.startswith(OLE2_SIGNATURE):
        if "spreadsheet" in lowered_type or "excel" in lowered_type or ext == ".xls":
            return ".xls"
        return ".doc" if ext not in {".xls"} else ext
    if "pdf" in lowered_type:
        return ".pdf"
    if "wordprocessingml" in lowered_type:
        return ".docx"
    if "spreadsheetml" in lowered_type:
        return ".xlsx"
    if "msword" in lowered_type:
        return ".doc"
    if "ms-excel" in lowered_type:
        return ".xls"
    return ext


def extract_attachment_text(content: bytes, url: str, name: str = "") -> str:
    return extract_attachment_payload(content, url, name).get("text", "")


def extract_attachment_payload(content: bytes, url: str, name: str = "", content_type: str = "") -> dict:
    ext = infer_attachment_extension(content, url, name, content_type)
    try:
        if ext == ".pdf":
            return _extract_pdf_payload(content)
        if ext == ".docx":
            return _extract_docx_payload(content)
        if ext == ".xlsx":
            return _extract_xlsx_payload(content)
        if ext in {".doc", ".xls"}:
            return _extract_legacy_office_payload(content)
    except Exception:
        return {"text": "", "pages": [], "sheets": [], "tables": []}
    return {"text": "", "pages": [], "sheets": [], "tables": []}


def render_attachment_context(attachments: list[dict[str, Any]]) -> str:
    parts: list[str] = []
    for item in attachments:
        name = item.get("name") or "附件"
        text = _clean_text(item.get("text") or "")
        if text:
            parts.append(f"附件《{name}》正文摘录：\n{text[:MAX_ATTACHMENT_TEXT_CHARS]}")
    return "\n\n".join(parts)


def attachment_text_parts(attachment: dict) -> list[dict[str, object]]:
    name = attachment.get("name") or "附件"
    parts: list[dict[str, object]] = []
    for item in attachment.get("pages") or []:
        text = _clean_text(item.get("text") or "")[:MAX_ATTACHMENT_PART_CHARS]
        if text:
            parts.append(
                {
                    "text": f"附件《{name}》第{item.get('page')}页：\n{text}",
                    "heading": f"附件：{name}",
                    "page": item.get("page"),
                    "attachment_name": name,
                }
            )
    for item in attachment.get("sheets") or []:
        text = _clean_text(item.get("text") or "")[:MAX_ATTACHMENT_PART_CHARS]
        if text:
            parts.append(
                {
                    "text": f"附件《{name}》工作表：{item.get('sheet')}\n{text}",
                    "heading": f"附件：{name} / {item.get('sheet')}",
                    "page": None,
                    "attachment_name": name,
                }
            )
    if not parts:
        text = _clean_text(attachment.get("text") or "")[:MAX_ATTACHMENT_PART_CHARS]
        if text:
            parts.append(
                {
                    "text": f"附件《{name}》正文摘录：\n{text}",
                    "heading": f"附件：{name}",
                    "page": None,
                    "attachment_name": name,
                }
            )
    return parts


def attachment_table_row_parts(attachment: dict) -> list[dict[str, object]]:
    name = attachment.get("name") or "附件"
    parts: list[dict[str, object]] = []
    for item in attachment.get("sheets") or []:
        label = str(item.get("sheet") or "工作表")
        rows = item.get("rows") if isinstance(item, dict) else None
        if not isinstance(rows, list) or not rows:
            rows = _rows_from_plain_table_text("sheet", label, str(item.get("text") or ""))
        parts.extend(_table_row_parts_from_rows(name, "工作表", label, rows))
        if len(parts) >= MAX_ATTACHMENT_TABLE_ROWS:
            return parts[:MAX_ATTACHMENT_TABLE_ROWS]
    for item in attachment.get("tables") or []:
        label = str(item.get("table") or "表格")
        rows = item.get("rows") if isinstance(item, dict) else None
        if not isinstance(rows, list) or not rows:
            rows = _rows_from_plain_table_text("table", label, str(item.get("text") or ""))
        parts.extend(_table_row_parts_from_rows(name, "表格", label, rows))
        if len(parts) >= MAX_ATTACHMENT_TABLE_ROWS:
            return parts[:MAX_ATTACHMENT_TABLE_ROWS]
    for item in attachment.get("pages") or []:
        if len(parts) >= MAX_PDF_TABLE_ROWS_PER_ATTACHMENT:
            break
        page_number = item.get("page")
        text = str(item.get("text") or "")
        if not _page_looks_table_like(name, text):
            continue
        rows = _rows_from_pdf_page_text(page_number, text)
        parts.extend(_table_row_parts_from_rows(name, "PDF页", f"第{page_number or '?'}页", rows))
        if len(parts) >= MAX_PDF_TABLE_ROWS_PER_ATTACHMENT:
            return parts[:MAX_PDF_TABLE_ROWS_PER_ATTACHMENT]
    return parts[:MAX_ATTACHMENT_TABLE_ROWS]


def _table_row_parts_from_rows(
    attachment_name: str,
    label_type: str,
    label: str,
    rows: list[dict[str, Any]],
) -> list[dict[str, object]]:
    parts: list[dict[str, object]] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        text = _clean_text(str(row.get("text") or ""))[:MAX_TABLE_ROW_TEXT_CHARS]
        if not text:
            continue
        row_number = row.get("row_number")
        heading = f"附件表格：{attachment_name} / {label_type}：{label}"
        if row_number:
            heading += f" / 第{row_number}行"
        parts.append(
            {
                "text": f"附件《{attachment_name}》{label_type}：{label}；第{row_number or '?'}行\n{text}",
                "heading": heading,
                "page": _page_number_from_label(label) if label_type == "PDF页" else None,
                "attachment_name": attachment_name,
                "sheet": label if label_type == "工作表" else None,
                "table": label if label_type == "表格" else None,
                "row_number": row_number,
            }
        )
    return parts


def _page_number_from_label(label: str) -> int | None:
    match = re.search(r"第(\d+)页", label or "")
    return int(match.group(1)) if match else None


def _page_looks_table_like(name: str, text: str) -> bool:
    haystack = f"{name}\n{text or ''}"
    strong_pattern = r"(一览表|汇总表|安排表|信息表|统计表|申请表|名单|清单|名册|接收.*表|各学院接收)"
    if re.search(strong_pattern, name or ""):
        return True
    head = (text or "")[:500]
    return bool(re.search(strong_pattern, head))


def _rows_from_pdf_page_text(page_number: Any, text: str) -> list[dict[str, Any]]:
    lines = [re.sub(r"\s+", " ", line).strip() for line in str(text or "").splitlines()]
    lines = [line for line in lines if 2 <= len(line) <= 260]
    if not lines:
        return []
    header_lines = lines[: min(5, len(lines))]
    rows: list[dict[str, Any]] = []
    keyword_pattern = re.compile(r"(学院|专业|名额|条件|要求|课程|考核|面试|笔试|成绩|备注|姓名|学号|时间|地点|材料)")
    for index, line in enumerate(lines, start=1):
        if len(rows) >= MAX_PDF_TABLE_ROWS_PER_PAGE:
            break
        if index <= len(header_lines) and len(lines) > len(header_lines):
            continue
        if not keyword_pattern.search(line) and len(line) < 12:
            continue
        window = lines[max(0, index - 3) : min(len(lines), index + 2)]
        context = "\n".join(dict.fromkeys([*header_lines, *window]))
        rows.append(
            {
                "row_number": index,
                "text": f"PDF第{page_number or '?'}页第{index}行；{context}",
                "cells": [{"column": "行文本", "value": line}],
            }
        )
    return rows


def _extract_pdf_payload(content: bytes) -> dict:
    try:
        import pypdf
    except Exception:
        return {"text": "", "pages": [], "sheets": [], "tables": []}
    reader = pypdf.PdfReader(io.BytesIO(content))
    pages = []
    for index, page in enumerate(reader.pages[:20], start=1):
        text = _clean_text(page.extract_text() or "")[:MAX_ATTACHMENT_PART_CHARS]
        if text:
            pages.append({"page": index, "text": text})
    return {
        "text": _clean_text("\n\n".join(f"第{item['page']}页：\n{item['text']}" for item in pages)),
        "pages": pages,
        "sheets": [],
        "tables": [],
    }


def _extract_docx_payload(content: bytes) -> dict:
    try:
        import docx
    except Exception:
        return _extract_docx_payload_with_zip(content)
    try:
        document = docx.Document(io.BytesIO(content))
        texts = [paragraph.text for paragraph in document.paragraphs]
        tables = []
        for table_index, table in enumerate(document.tables[:20], start=1):
            raw_rows = []
            for row in table.rows:
                if len(raw_rows) >= 120:
                    break
                raw_rows.append((len(raw_rows) + 1, [_clean_cell(cell.text) for cell in row.cells]))
            table_payload = _table_payload("table", f"表格{table_index}", raw_rows)
            if table_payload:
                tables.append(table_payload)
                texts.append(f"表格{table_index}\n{table_payload['text']}")
        text = _clean_text("\n".join(texts))
    except Exception:
        return _extract_docx_payload_with_zip(content)
    if not text:
        return _extract_docx_payload_with_zip(content)
    return {"text": text, "pages": [], "sheets": [], "tables": tables}


def _extract_xlsx_payload(content: bytes) -> dict:
    try:
        import openpyxl
    except Exception:
        return _extract_xlsx_payload_with_zip(content)
    try:
        workbook = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        sheets = []
        for sheet in workbook.worksheets[:5]:
            raw_rows = []
            for row_number, row in enumerate(sheet.iter_rows(max_row=120, values_only=True), start=1):
                raw_rows.append((row_number, [_clean_cell(value) for value in row]))
            sheet_payload = _table_payload("sheet", sheet.title, raw_rows)
            if sheet_payload:
                sheet_payload["text"] = sheet_payload["text"][:MAX_ATTACHMENT_PART_CHARS]
                sheets.append(sheet_payload)
    except Exception:
        return _extract_xlsx_payload_with_zip(content)
    if not sheets:
        return _extract_xlsx_payload_with_zip(content)
    return {
        "text": _clean_text("\n\n".join(f"工作表：{item['sheet']}\n{item['text']}" for item in sheets)),
        "pages": [],
        "sheets": sheets,
        "tables": [],
    }


def _extract_docx_payload_with_zip(content: bytes) -> dict:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            document_xml = archive.read("word/document.xml")
    except Exception:
        return {"text": "", "pages": [], "sheets": [], "tables": []}
    try:
        root = ET.fromstring(document_xml)
    except ET.ParseError:
        return {"text": "", "pages": [], "sheets": [], "tables": []}
    texts = []
    for paragraph in root.iter(f"{WORD_NS}p"):
        paragraph_text = "".join(node.text or "" for node in paragraph.iter(f"{WORD_NS}t")).strip()
        if paragraph_text:
            texts.append(paragraph_text)
    tables = []
    for table_index, table in enumerate(root.iter(f"{WORD_NS}tbl"), start=1):
        if table_index > 20:
            break
        raw_rows = []
        for row_number, row in enumerate(table.iter(f"{WORD_NS}tr"), start=1):
            if row_number > 120:
                break
            cells = []
            for cell in row.findall(f"{WORD_NS}tc"):
                cell_text = " ".join(
                    "".join(node.text or "" for node in paragraph.iter(f"{WORD_NS}t")).strip()
                    for paragraph in cell.iter(f"{WORD_NS}p")
                )
                cells.append(_clean_cell(cell_text))
            raw_rows.append((row_number, cells))
        table_payload = _table_payload("table", f"表格{table_index}", raw_rows)
        if table_payload:
            tables.append(table_payload)
            texts.append(f"表格{table_index}\n{table_payload['text']}")
    text = _clean_text("\n".join(texts))
    return {"text": text, "pages": [], "sheets": [], "tables": tables}


def _extract_xlsx_payload_with_zip(content: bytes) -> dict:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            shared_strings = _read_shared_strings(archive)
            sheets = []
            for sheet_name, sheet_path in _xlsx_sheet_paths(archive)[:5]:
                if sheet_path not in archive.namelist():
                    continue
                sheet_payload = _read_sheet_payload(sheet_name, archive.read(sheet_path), shared_strings)
                if sheet_payload:
                    sheet_payload["text"] = sheet_payload["text"][:MAX_ATTACHMENT_PART_CHARS]
                    sheets.append(sheet_payload)
    except Exception:
        return {"text": "", "pages": [], "sheets": [], "tables": []}
    return {
        "text": _clean_text("\n\n".join(f"工作表：{item['sheet']}\n{item['text']}" for item in sheets)),
        "pages": [],
        "sheets": sheets,
        "tables": [],
    }


def _read_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    if "xl/sharedStrings.xml" not in archive.namelist():
        return []
    try:
        root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
    except ET.ParseError:
        return []
    values = []
    for item in root.iter(f"{SHEET_NS}si"):
        text = "".join(node.text or "" for node in item.iter(f"{SHEET_NS}t")).strip()
        values.append(text)
    return values


def _xlsx_sheet_paths(archive: zipfile.ZipFile) -> list[tuple[str, str]]:
    names = archive.namelist()
    if "xl/workbook.xml" not in names:
        return [(Path(path).stem, path) for path in sorted(names) if path.startswith("xl/worksheets/")]

    rels: dict[str, str] = {}
    if "xl/_rels/workbook.xml.rels" in names:
        try:
            rel_root = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
            for rel in rel_root.iter(f"{REL_NS}Relationship"):
                rel_id = rel.attrib.get("Id")
                target = rel.attrib.get("Target", "")
                if rel_id and target:
                    target = target.lstrip("/")
                    rels[rel_id] = target if target.startswith("xl/") else f"xl/{target}"
        except ET.ParseError:
            rels = {}

    paths: list[tuple[str, str]] = []
    try:
        root = ET.fromstring(archive.read("xl/workbook.xml"))
    except ET.ParseError:
        return []
    for index, sheet in enumerate(root.iter(f"{SHEET_NS}sheet"), start=1):
        name = sheet.attrib.get("name") or f"sheet{index}"
        rel_id = sheet.attrib.get(f"{OFFICE_REL_NS}id")
        path = rels.get(rel_id or "", f"xl/worksheets/sheet{index}.xml")
        paths.append((name, path))
    return paths


def _read_sheet_payload(sheet_name: str, sheet_xml: bytes, shared_strings: list[str]) -> dict[str, Any] | None:
    try:
        root = ET.fromstring(sheet_xml)
    except ET.ParseError:
        return None
    raw_rows = []
    for row_index, row in enumerate(root.iter(f"{SHEET_NS}row"), start=1):
        if row_index > 120:
            break
        values: list[str] = []
        for cell in row.iter(f"{SHEET_NS}c"):
            value = _cell_text(cell, shared_strings)
            values.append(value)
        raw_rows.append((int(row.attrib.get("r") or row_index), values))
    return _table_payload("sheet", sheet_name, raw_rows)


def _read_sheet_text(sheet_xml: bytes, shared_strings: list[str]) -> str:
    payload = _read_sheet_payload("Sheet", sheet_xml, shared_strings)
    return str(payload.get("text") or "")[:MAX_ATTACHMENT_PART_CHARS] if payload else ""


def _cell_text(cell: ET.Element, shared_strings: list[str]) -> str:
    cell_type = cell.attrib.get("t")
    if cell_type == "inlineStr":
        value = "".join(node.text or "" for node in cell.iter(f"{SHEET_NS}t"))
        return value.strip()
    value_node = cell.find(f"{SHEET_NS}v")
    if value_node is None or value_node.text is None:
        return ""
    raw = value_node.text.strip()
    if cell_type == "s":
        try:
            return shared_strings[int(raw)].strip()
        except (ValueError, IndexError):
            return ""
    return raw


def _extract_legacy_office_payload(content: bytes) -> dict:
    text = _extract_readable_binary_text(content)
    return {"text": text, "pages": [], "sheets": [], "tables": []}


def _table_payload(label_key: str, label: str, raw_rows: list[tuple[int, list[Any]]]) -> dict[str, Any] | None:
    cleaned_rows = [
        (row_number, _trim_trailing_empty_cells([_clean_cell(value) for value in cells]))
        for row_number, cells in raw_rows
    ]
    cleaned_rows = [(row_number, cells) for row_number, cells in cleaned_rows if any(cells)]
    if not cleaned_rows:
        return None
    lines = [" | ".join(value for value in cells if value) for _, cells in cleaned_rows]
    header_number, headers = _table_headers(cleaned_rows)
    rows = []
    for row_number, cells in cleaned_rows:
        if row_number == header_number and len(cleaned_rows) > 1:
            continue
        row_text, cell_pairs = _format_table_row(label_key, label, row_number, headers, cells)
        if row_text:
            rows.append({"row_number": row_number, "text": row_text, "cells": cell_pairs})
    payload = {
        label_key: label,
        "text": _clean_text("\n".join(line for line in lines if line)),
        "rows": rows[:MAX_ATTACHMENT_TABLE_ROWS],
    }
    return payload if payload["text"] else None


def _table_headers(cleaned_rows: list[tuple[int, list[str]]]) -> tuple[int | None, list[str]]:
    if len(cleaned_rows) <= 1:
        max_len = max((len(cells) for _, cells in cleaned_rows), default=0)
        return None, [f"列{index + 1}" for index in range(max_len)]
    row_number, cells = cleaned_rows[0]
    headers = []
    for index, value in enumerate(cells):
        header = re.sub(r"\s+", "", value or "")
        headers.append(header or f"列{index + 1}")
    return row_number, headers


def _format_table_row(
    label_key: str,
    label: str,
    row_number: int,
    headers: list[str],
    cells: list[str],
) -> tuple[str, list[dict[str, str]]]:
    pairs = []
    for index, value in enumerate(cells):
        if not value:
            continue
        column = headers[index] if index < len(headers) and headers[index] else f"列{index + 1}"
        pairs.append({"column": column, "value": value})
    if not pairs:
        return "", []
    label_type = "工作表" if label_key == "sheet" else "表格"
    facts = "；".join(f"{item['column']}：{item['value']}" for item in pairs)
    text = f"{label_type}：{label}；第{row_number}行；{facts}"
    return text[:MAX_TABLE_ROW_TEXT_CHARS], pairs


def _rows_from_plain_table_text(label_key: str, label: str, text: str) -> list[dict[str, Any]]:
    raw_rows = []
    for row_number, line in enumerate(str(text or "").splitlines(), start=1):
        line = line.strip()
        if not line:
            continue
        cells = [item.strip() for item in re.split(r"\s*\|\s*|\t+", line)]
        raw_rows.append((row_number, cells))
    payload = _table_payload(label_key, label, raw_rows)
    return list(payload.get("rows") or []) if payload else []


def _clean_cell(value: Any) -> str:
    if value is None:
        return ""
    text = re.sub(r"\s+", " ", str(value)).strip()
    return text[:180]


def _trim_trailing_empty_cells(cells: list[str]) -> list[str]:
    output = list(cells)
    while output and not output[-1]:
        output.pop()
    return output


def _extract_readable_binary_text(content: bytes) -> str:
    candidates = []
    for encoding in ("utf-16le", "gb18030", "utf-8", "latin1"):
        try:
            decoded = content.decode(encoding, errors="ignore")
        except Exception:
            continue
        candidates.append(_legacy_text_from_decoded(decoded))
    best = max(candidates, key=_legacy_text_score, default="")
    if _legacy_text_score(best) < MIN_LEGACY_TEXT_SCORE:
        return ""
    return _clean_text(best)


def _legacy_text_from_decoded(text: str) -> str:
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]+", "\n", text)
    fragments = re.findall(r"[\u4e00-\u9fffA-Za-z0-9，。；：、（）()《》“”\-_/ ]{2,}", text)
    lines = []
    current = ""
    for fragment in fragments:
        fragment = _clean_legacy_fragment(fragment)
        if _legacy_fragment_score(fragment) <= 0:
            continue
        if len(current) + len(fragment) > 120:
            if current:
                lines.append(current)
            current = fragment
        else:
            current = f"{current} {fragment}".strip() if current else fragment
        if len("\n".join(lines)) > MAX_ATTACHMENT_TEXT_CHARS:
            break
    if current:
        lines.append(current)
    return "\n".join(lines)


def _clean_legacy_fragment(text: str) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    text = re.sub(r"\bPAGEREF\s+_Toc\d+\b", " ", text)
    text = re.sub(r"\b_Toc\d+\b", " ", text)
    text = re.sub(r"\b[A-Za-z0-9+/=_-]{24,}\b", " ", text)
    text = re.sub(r"\b(?:PK|xml|rels)\b", " ", text)
    for term in LEGACY_OFFICE_NOISE_TERMS:
        text = re.sub(re.escape(term), " ", text, flags=re.IGNORECASE)
    text = re.sub(r"(?:theme/){1,}\w*", " ", text, flags=re.IGNORECASE)
    return re.sub(r"\s+", " ", text).strip(" -_/")


def _legacy_text_score(text: str) -> int:
    return sum(_legacy_fragment_score(line) for line in (text or "").splitlines())


def _legacy_fragment_score(text: str) -> int:
    text = (text or "").strip()
    if len(text) < 2:
        return 0
    noise_hits = sum(1 for term in LEGACY_OFFICE_NOISE_TERMS if term.lower() in text.lower())
    common_hits = sum(1 for term in LEGACY_COMMON_TERMS if term in text)
    cjk = len(re.findall(r"[\u4e00-\u9fff]", text))
    english_words = [
        word
        for word in re.findall(r"[A-Za-z]{3,}", text)
        if word.lower() not in {"root", "entry", "summaryinformation", "documentsummaryinformation", "worddocument"}
    ]
    english_chars = sum(len(word) for word in english_words)
    if noise_hits >= 2 and common_hits == 0:
        return 0
    if common_hits:
        return cjk * 2 + english_chars + common_hits * 10
    if cjk >= 12 and _common_cjk_ratio(text) >= 0.35:
        return cjk
    if cjk >= 2 and len(english_words) >= 2 and noise_hits == 0:
        return english_chars
    return 0


def _common_cjk_ratio(text: str) -> float:
    chars = re.findall(r"[\u4e00-\u9fff]", text or "")
    if not chars:
        return 0
    common_chars = set("的一是在不了有和人这中大为上个国我以要他时来用们生到作地于出就分对成会可主发年动同工也能下过子说产种面而方后多定行学法所民得经十三之进着等部度家电力里如水化高自二理起小物现实加量都两体制机当使点从业本去把性好应开它合还因由其些然前外天政四日那社义事平形相全表间样与关各重新线内数正心反你明看原又么利比或但质气第向道命此变条只没结解问意建月公无系军很情者最立代想已通并提直题党程展五果料象员革位入常文总次品式活设及管特件长求老头基资边流路级少图山统接知较将组见计别她手角期根论运农指几九区强放决西被干做必战先回则任取据处队南给色光门即保治北造百规热领七海口东导器压志世金增争济阶油思术极交受联什认六共权收证改清己美再采转更单风切打白教速花带安场身车例真务具万每目至达走积示议声报斗完类八离华名确才科张信马节话米整空元况今集温传土许步群广石记需段研界拉林律叫且究观越织装影算低持音众书布复容儿须际商非验连断深难近矿千周委素技备半办青省列习响约支般史感劳便团往酸历市克何除消构府称太准精值号率族维划选标写存候毛亲快效斯院查江型眼王按格养易置派层片始却专状育厂京识适属圆包火住调满县局照参红细引听该铁价严龙飞")
    return sum(1 for char in chars if char in common_chars) / len(chars)


def _text_quality(text: str) -> int:
    cjk = len(re.findall(r"[\u4e00-\u9fff]", text or ""))
    ascii_chars = len(re.findall(r"[A-Za-z0-9]", text or ""))
    return cjk * 2 + ascii_chars


def _zip_has(content: bytes, member: str) -> bool:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            return member in archive.namelist()
    except Exception:
        return False


def _clean_text(text: str) -> str:
    text = re.sub(r"[ \t]{2,}", " ", text or "")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:MAX_ATTACHMENT_TEXT_CHARS]
